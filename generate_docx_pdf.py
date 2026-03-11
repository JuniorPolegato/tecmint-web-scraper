#!/usr/bin/env python3
"""
Generate a .docx from table.md that looks like the GitHub preview,
then convert to PDF using LibreOffice.
"""

import os
import re
import subprocess
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


IMAGES_DIR = "tool_images"
MD_FILE = "table.md"
DOCX_FILE = "60_Best_Free_Open_Source_Software.docx"
PDF_FILE = "60_Best_Free_Open_Source_Software.pdf"

# Page config
PAGE_MARGIN = Cm(1.5)
IMG_SIZE = Inches(0.55)


def convert_webp_to_png(path):
    """Convert .webp to .png in memory using Pillow, return bytes."""
    img = Image.open(path)
    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def get_image_stream(img_path):
    """Return a file-like object for the image, converting ALL through Pillow."""
    if not os.path.exists(img_path):
        return None
    ext = os.path.splitext(img_path)[1].lower()
    if ext == ".svg":
        return None  # docx doesn't support SVG
    try:
        # Force convert everything through Pillow to PNG
        img = Image.open(img_path)
        img = img.convert("RGBA")
        buf = BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf
    except Exception as e:
        print(f"  ⚠️  Pillow can't open: {img_path} — {e}")
        return None


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=50, bottom=50, left=80, right=80):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:start w:w="{left}" w:type="dxa"/>'
        f'  <w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def set_row_height(row, height_cm):
    """Set exact row height."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{int(height_cm * 567)}" w:hRule="atLeast"/>'
    )
    trPr.append(trHeight)


def remove_paragraph_spacing(paragraph):
    """Remove extra spacing from paragraph."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(12)


def parse_table_md(md_file):
    """Parse table.md and extract structure: headings, categories, tools."""
    with open(md_file, "r", encoding="utf-8") as f:
        lines = f.readlines()

    elements = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Main title: # ...
        if line.startswith("# ") and not line.startswith("## "):
            elements.append({"type": "title", "text": line[2:].strip()})

        # Blockquote (source)
        elif line.startswith("> "):
            elements.append({"type": "quote" if not elements or elements[-1]["type"] != "tool_heading" else "description", "text": line[2:].strip()})

        # Category heading: ## 📁 ...
        elif line.startswith("## "):
            elements.append({"type": "category", "text": line[3:].strip()})

        # Tool heading: ### N. Tool Name
        elif line.startswith("### "):
            elements.append({"type": "tool_heading", "text": line[4:].strip()})

        # Links line: 🔗 [text](url) | ...
        elif line.startswith("🔗") or line.startswith("📎"):
            links = re.findall(r'\[([^\]]+)\]\(([^)]+)\)', line)
            elements.append({"type": "links", "links": links})

        # Image: ![alt](path) — handles parentheses in filenames
        elif line.startswith("!["):
            match = re.match(r'!\[([^\]]*)\]\((.+)\)$', line)
            if match:
                elements.append({"type": "image", "alt": match.group(1), "path": match.group(2)})

        # Table row: | ... | ... | ... |
        elif line.startswith("|") and "---" not in line:
            cols = [c.strip() for c in line.split("|")[1:-1]]
            if len(cols) >= 3:
                if cols[0] in ("#", "Nº", "No", "No."):
                    elements.append({"type": "table_header", "cols": cols})
                else:
                    elements.append({"type": "table_row", "cols": cols})

        # Horizontal rule
        elif line.startswith("---"):
            elements.append({"type": "hr"})

        i += 1

    return elements


def build_docx(elements):
    """Build the .docx document."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = PAGE_MARGIN
        section.bottom_margin = PAGE_MARGIN
        section.left_margin = PAGE_MARGIN
        section.right_margin = PAGE_MARGIN

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)

    table_rows = []
    table_header = None

    def flush_table():
        nonlocal table_rows, table_header
        if not table_header or not table_rows:
            table_header = None
            table_rows = []
            return

        ncols = len(table_header["cols"])
        table = doc.add_table(rows=1 + len(table_rows), cols=ncols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # Column widths
        col_widths = [Cm(1), Cm(2.5), Cm(4), Cm(10.5)]
        for idx, width in enumerate(col_widths[:ncols]):
            for row in table.rows:
                row.cells[idx].width = width

        # Header row
        hdr_row = table.rows[0]
        set_row_height(hdr_row, 0.7)
        for j, col_text in enumerate(table_header["cols"]):
            cell = hdr_row.cells[j]
            set_cell_shading(cell, "24292E")
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            remove_paragraph_spacing(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(col_text)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Data rows
        for i, row_data in enumerate(table_rows):
            row = table.rows[i + 1]
            set_row_height(row, 1.2)
            bg = "F6F8FA" if i % 2 == 0 else "FFFFFF"

            for j, col_text in enumerate(row_data["cols"][:ncols]):
                cell = row.cells[j]
                set_cell_shading(cell, bg)
                set_cell_margins(cell)
                p = cell.paragraphs[0]
                remove_paragraph_spacing(p)

                # Column 0: number (centered)
                if j == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(col_text)
                    run.font.size = Pt(9)
                    run.bold = True

                # Column 1: image
                elif j == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_match = re.match(r'!\[([^\]]*)\]\((.+)\)$', col_text.strip())
                    if img_match:
                        img_path = img_match.group(2)
                        stream = get_image_stream(img_path)
                        if stream:
                            try:
                                run = p.add_run()
                                run.add_picture(stream, width=IMG_SIZE)
                            except Exception as e:
                                run = p.add_run("🖼️")
                                print(f"  ⚠️  Image error: {img_path} — {e}")
                            finally:
                                if hasattr(stream, 'close'):
                                    stream.close()
                        else:
                            run = p.add_run("🖼️")
                    else:
                        run = p.add_run(col_text)
                        run.font.size = Pt(9)

                # Column 2: name (bold)
                elif j == 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    bold_match = re.match(r'\*\*(.+?)\*\*', col_text)
                    if bold_match:
                        run = p.add_run(bold_match.group(1))
                        run.bold = True
                    else:
                        run = p.add_run(col_text)
                    run.font.size = Pt(9)

                # Column 3: description
                elif j == 3:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(col_text)
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0x58, 0x60, 0x69)

        # Table borders
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D0D7DE"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D7DE"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D0D7DE"/>'
            f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="D0D7DE"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

        table_header = None
        table_rows = []

    # Process elements
    for el in elements:
        t = el["type"]

        if t == "title":
            flush_table()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(el["text"])
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)

        elif t == "quote":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            text = el["text"]
            parts = re.split(r'(\[[^\]]+\]\([^)]+\))', text)
            for part in parts:
                link_match = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
                if link_match:
                    run = p.add_run(link_match.group(1))
                    run.font.color.rgb = RGBColor(0x08, 0x69, 0xDA)
                    run.font.size = Pt(9)
                    run.italic = True
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(9)
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x58, 0x60, 0x69)

        elif t == "category":
            flush_table()
            p = doc.add_heading(level=2)
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(el["text"])
            run.font.name = "Calibri"
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)

        elif t == "tool_heading":
            flush_table()
            p = doc.add_heading(level=3)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(el["text"])
            run.font.name = "Calibri"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)

        elif t == "description":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="D0D7DE"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="284"/>')
            pPr.append(ind)
            run = p.add_run(el["text"])
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x58, 0x60, 0x69)
            run.italic = True

        elif t == "links":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run("🔗 ")
            run.font.size = Pt(9)
            for idx, (text, url) in enumerate(el["links"]):
                if idx > 0:
                    run = p.add_run(" | ")
                    run.font.size = Pt(9)
                run = p.add_run(text)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x08, 0x69, 0xDA)
                run.underline = True

            # If next element is NOT an image, add page break here
            el_index = elements.index(el)
            next_is_image = (
                el_index + 1 < len(elements)
                and elements[el_index + 1]["type"] == "image"
            )
            is_last_tool = all(
                elements[k]["type"] in ("hr", "table_header", "table_row")
                for k in range(el_index + 1, len(elements))
            )
            if not next_is_image and not is_last_tool:
                p_break = doc.add_paragraph()
                run_break = p_break.add_run()
                run_break.add_break(WD_BREAK.PAGE)
                remove_paragraph_spacing(p_break)

        elif t == "image":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            stream = get_image_stream(el["path"])
            if stream:
                try:
                    run = p.add_run()
                    run.add_picture(stream, width=Inches(5.5))
                except Exception as e:
                    run = p.add_run(f"[Image: {el['alt']}]")
                    print(f"  ⚠️  Image error: {el['path']} — {e}")
                finally:
                    if hasattr(stream, 'close'):
                        stream.close()

            # Page break after image (one tool per page)
            # But skip if this is the last element to avoid blank page
            el_index = elements.index(el)
            if el_index < len(elements) - 1:
                p_break = doc.add_paragraph()
                run_break = p_break.add_run()
                run_break.add_break(WD_BREAK.PAGE)
                remove_paragraph_spacing(p_break)

        elif t == "table_header":
            table_header = el

        elif t == "table_row":
            table_rows.append(el)

        elif t == "hr":
            flush_table()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="D0D7DE"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

    # Flush any remaining table
    flush_table()

    return doc


def main():
    print("=" * 60)
    print("📄 GENERATING DOCX FROM table.md")
    print("=" * 60)

    # Parse
    print(f"\n📖 Parsing {MD_FILE}...")
    elements = parse_table_md(MD_FILE)
    print(f"   Found {len(elements)} elements")

    # Build docx
    print(f"\n🔨 Building {DOCX_FILE}...")
    doc = build_docx(elements)
    doc.save(DOCX_FILE)
    print(f"   ✅ Saved: {DOCX_FILE}")

    # Convert to PDF with LibreOffice
    print(f"\n📄 Converting to PDF with LibreOffice...")
    try:
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", ".",
                DOCX_FILE
            ],
            capture_output=True,
            text=True,
            timeout=120
        )
        if result.returncode == 0:
            print(f"   ✅ Saved: {PDF_FILE}")
        else:
            print(f"   ❌ LibreOffice error: {result.stderr}")
    except FileNotFoundError:
        print("   ❌ LibreOffice not found! Install with:")
        print("      sudo apt install libreoffice")
    except subprocess.TimeoutExpired:
        print("   ❌ LibreOffice timed out (120s)")

    # Final report
    print(f"\n{'=' * 60}")
    print(f"📊 DONE!")
    print(f"{'=' * 60}")
    if os.path.exists(DOCX_FILE):
        size_kb = os.path.getsize(DOCX_FILE) / 1024
        print(f"  📝 DOCX: {DOCX_FILE} ({size_kb:.0f} KB)")
    if os.path.exists(PDF_FILE):
        size_kb = os.path.getsize(PDF_FILE) / 1024
        print(f"  📄 PDF:  {PDF_FILE} ({size_kb:.0f} KB)")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
